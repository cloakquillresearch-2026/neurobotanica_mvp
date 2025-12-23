"""Build a Word (.docx) filing-ready provisional with embedded figures.

Source: docs/Provisional Patent Application NeuroBotanica.sty
Figures: docs/figures/Figure_0X_*.png
Output: docs/NeuroBotanica_Provisional_Patent_Application.docx

This script keeps formatting conservative (headings, bold, bullets) and
inserts FIG. 1–FIG. 8 images with captions after the "BRIEF DESCRIPTION OF THE DRAWINGS" section.
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


FIG_RE = re.compile(r"\*\*Figure\s+(\d+)\*\*\s+(.*)")


@dataclass
class FigureDesc:
    num: int
    desc: str


def add_runs_with_bold(paragraph, text: str) -> None:
    """Add text to paragraph, interpreting **bold** spans."""
    if "**" not in text:
        paragraph.add_run(text)
        return

    parts = text.split("**")
    for i, part in enumerate(parts):
        run = paragraph.add_run(part)
        if i % 2 == 1:
            run.bold = True


def normalize_caption(desc: str) -> str:
    d = desc.strip()
    lower = d.lower()
    for prefix in ("illustrates ", "depicts ", "shows ", "details "):
        if lower.startswith(prefix):
            d = d[len(prefix) :].lstrip()
            break
    if d.lower().startswith("the "):
        d = d[4:]
    if d:
        d = d[0].upper() + d[1:]
    return d


def is_heading(line: str) -> tuple[int, str] | None:
    if line.startswith("### "):
        return 2, line[4:].strip()
    if line.startswith("## "):
        return 1, line[3:].strip()
    return None


def build_docx(src_path: Path, figures_dir: Path, out_path: Path) -> None:
    text = src_path.read_text(encoding="utf-8", errors="replace")
    lines = [ln.rstrip("\r\n") for ln in text.splitlines()]

    doc = Document()

    # Slightly wider layout by default; leave margins as Word defaults.

    in_bullets = False
    brief_started = False
    injected_figures = False
    fig_descs: dict[int, str] = {}

    def flush_bullet_mode():
        nonlocal in_bullets
        in_bullets = False

    for line in lines:
        stripped = line.strip()

        # Collect figure descriptions for captions
        m = FIG_RE.search(line)
        if m:
            num = int(m.group(1))
            desc = m.group(2).strip()
            fig_descs[num] = desc

        heading = is_heading(stripped)
        if heading is not None:
            flush_bullet_mode()
            level, title = heading
            doc.add_heading(title, level=level)
            if title.upper() == "BRIEF DESCRIPTION OF THE DRAWINGS":
                brief_started = True
            continue

        if stripped == "---":
            flush_bullet_mode()

            # After the brief description block, insert the actual figures once.
            if brief_started and not injected_figures:
                # The brief description ends at the first --- following it.
                doc.add_paragraph("")
                for i in range(1, 9):
                    png = figures_dir / f"Figure_{i:02d}_"  # prefix
                    matches = sorted(figures_dir.glob(png.name + "*.png"))
                    if not matches:
                        raise FileNotFoundError(f"Missing PNG for Figure {i} (expected Figure_{i:02d}_*.png)")

                    caption_desc = normalize_caption(fig_descs.get(i, ""))
                    caption = f"FIG. {i}. {caption_desc}".strip()

                    cap_p = doc.add_paragraph()
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = cap_p.add_run(caption)
                    cap_run.bold = True

                    pic_p = doc.add_paragraph()
                    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pic_p.add_run().add_picture(str(matches[0]), width=Inches(6.5))

                    doc.add_paragraph("")

                injected_figures = True

            # Keep section separation
            doc.add_paragraph("")
            continue

        if stripped.startswith("- "):
            in_bullets = True
            p = doc.add_paragraph(style="List Bullet")
            add_runs_with_bold(p, stripped[2:])
            continue

        if stripped == "":
            flush_bullet_mode()
            doc.add_paragraph("")
            continue

        # Title-ish early lines (bold markdown only)
        if stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            flush_bullet_mode()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_runs_with_bold(p, stripped)
            continue

        flush_bullet_mode()
        p = doc.add_paragraph()
        add_runs_with_bold(p, line)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> None:
    repo_root = Path(__file__).resolve().parents[1]
    src_path = repo_root / "docs" / "Provisional Patent Application NeuroBotanica.sty"
    figures_dir = repo_root / "docs" / "figures"
    out_path = repo_root / "docs" / "NeuroBotanica_Provisional_Patent_Application.docx"

    build_docx(src_path=src_path, figures_dir=figures_dir, out_path=out_path)
    print(f"Wrote: {out_path}")


if __name__ == "__main__":
    main()
